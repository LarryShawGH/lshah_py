import imaplib
import email
from email.header import decode_header
import docx
import os

# Your Yahoo email credentials
EMAIL = "lalit.shah@yahoo.com"
PASSWORD = "dxiforjqbmnhrrfd"
SENDER = "thenuancedperspective+10-day-agentic-ai-course@substack.com"

# Create a Word document
doc = docx.Document()
doc.add_heading("Emails from " + SENDER, 0)

def read_yahoo_emails():
    try:
        # Connect to Yahoo's IMAP server
        imap = imaplib.IMAP4_SSL("imap.mail.yahoo.com")

        # Login to your account
        imap.login(EMAIL, PASSWORD)

        # Select the mailbox you want to use (e.g., "INBOX")
        imap.select("INBOX")

        # Search for all emails in the mailbox
        status, messages = imap.search(None, "ALL")

        status, messages = imap.search(None, f'FROM "{SENDER}"')

        # Convert messages to a list of email IDs
        email_ids = messages[0].split()

        # Fetch the latest email (you can loop through all emails if needed)
        for email_id in email_ids[-2:]:  # Fetch the last 5 emails
            # Fetch the email by ID
            status, msg_data = imap.fetch(email_id, "(RFC822)")

            for response_part in msg_data:
                if isinstance(response_part, tuple):
                    # Parse the raw email bytes to a message object
                    msg = email.message_from_bytes(response_part[1])

                    # Decode the email subject
                    subject, encoding = decode_header(msg["Subject"])[0]
                    if isinstance(subject, bytes):
                        # If it's a bytes type, decode to str
                        subject = subject.decode(encoding if encoding else "utf-8")
                    print("Subject:", subject)

                    # Decode the sender's email address
                    from_ = msg.get("From")
                    # print("From:", from_)

                    # If the email has a body, print it
                    if msg.is_multipart():
                        for part in msg.walk():
                            # If the content type is text/plain or text/html
                            if part.get_content_type() == "text/plain":
                                body = part.get_payload(decode=True).decode()
                                print("Body:", body)
                    
                                break
                    else:
                        # If the email is not multipart
                        body = msg.get_payload(decode=True).decode()
                        print("Body:", body)

                    # Add to document
                    # doc.add_paragraph(f"Date: {date}")
                    doc.add_paragraph(f"Subject: {subject}")
                    doc.add_paragraph(f"Body: {body}")
                    doc.add_paragraph(f"End of email\n\n")


        # Close the connection and logout
        imap.close()
        imap.logout()
        
        # Set the path where you want to save the document
        # save_path = r'C:\Users\lshah\Documents'  # Make sure this folder exists
        # file_name = 'email.docx'
        # full_path = os.path.join(save_path, file_name)

        # Save the document
        # print(f'Document will be saved to: {full_path}')
        # doc.save(full_path)

        
        # Save the document
        doc.save("email_summary.docx")
        print(f'Document saved as email_summary.docx')

    except Exception as e:
        print("Error:", str(e))

if __name__ == "__main__":
    read_yahoo_emails()